#!/usr/bin/env python
# -*- coding: utf-8 -*-

import sys, time
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import Font
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

import csv
from collections import OrderedDict

csv.dict = OrderedDict
reader = csv.DictReader(open(sys.argv[1]))

document = docx.Document()
style = document.styles.add_style("hebrew", WD_STYLE_TYPE.PARAGRAPH)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font = style.font
font.complex_script = True
font.name = "David"
font.rtl = True 
font.size = Pt(12)

for line in reader:
    i = line.iteritems()
    for k,v in i:
        p = document.add_paragraph(style="hebrew")
        r = p.add_run(k.decode("utf8") + ": ")
        r.font.cs_bold = True
        r.font.rtl = True
        r = p.add_run(v.decode("utf8"))
        r.font.rtl = True
    document.add_page_break()

document.save("rasham_%s.docx" % time.strftime("%Y%m%d_%H%M%S"))
